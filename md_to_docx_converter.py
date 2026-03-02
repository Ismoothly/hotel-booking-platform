#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX Converter with full editing support
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def md_to_docx(markdown_file, output_file):
    """Convert Markdown file to DOCX"""
    
    # Read markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    in_comment = False
    
    while i < len(lines):
        line = lines[i]
        
        # Handle HTML comments (ignore them)
        if '<!--' in line:
            in_comment = True
        if '-->' in line:
            in_comment = False
            i += 1
            continue
        if in_comment:
            i += 1
            continue
        
        # Skip front matter
        if line.strip().startswith('---') and i == 0:
            # Skip until second ---
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('---'):
                i += 1
            i += 1
            continue
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_lines = []
            else:
                # End of code block
                in_code_block = False
                if code_block_lines:
                    # Add code block to document
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    for j, code_line in enumerate(code_block_lines):
                        if j > 0:
                            p.add_run('\n')
                        run = p.add_run(code_line)
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    # Add gray background
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'E7E6E6')
                    p._element.get_or_add_pPr().append(shading_elm)
            i += 1
            continue
        elif in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('# ').strip()
            
            if title:
                # Limit heading level to 9 (Word supports up to 9)
                heading_level = min(level, 9)
                p = doc.add_heading(title, level=heading_level)
                # Make heading bold and larger
                for run in p.runs:
                    run.font.size = Pt(14 + (10 - heading_level) * 1.5)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Handle tables
        if line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            # Create table
            if len(table_lines) >= 3:
                rows = []
                for tline in table_lines:
                    cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                    if cells:
                        rows.append(cells)
                
                if len(rows) >= 2:
                    # Filter out separator row
                    header_row = rows[0]
                    data_rows = [r for j, r in enumerate(rows[1:]) if not all('-' in c for c in r)]
                    
                    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_row))
                    table.style = 'Light Grid Accent 1'
                    
                    # Add header
                    for j, cell_text in enumerate(header_row):
                        cell = table.rows[0].cells[j]
                        cell.text = cell_text
                        # Make header bold
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Add data rows
                    for row_idx, row_data in enumerate(data_rows, 1):
                        for j, cell_text in enumerate(row_data):
                            if j < len(table.rows[row_idx].cells):
                                table.rows[row_idx].cells[j].text = cell_text
            continue
        
        # Handle lists
        if line.lstrip().startswith(('- ', '* ', '+ ')) or re.match(r'^\s*\d+\.\s', line):
            # Get indent level
            indent = len(line) - len(line.lstrip())
            
            # Extract list marker and content
            if line.lstrip().startswith(('- ', '* ', '+ ')):
                content = line.lstrip()[2:].strip()
            else:
                match = re.match(r'\s*\d+\.\s(.+)', line)
                content = match.group(1) if match else line.strip()
            
            if content:
                # Calculate indent level (0, 1, 2, etc.)
                level = max(0, indent // 2)
                p = doc.add_paragraph(content, style='List Bullet' if '- ' in line or '* ' in line or '+ ' in line else 'List Number')
                p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
            
            i += 1
            continue
        
        # Handle emphasis and bold
        if line.strip():
            p = doc.add_paragraph()
            
            # Process inline formatting
            remaining = line
            while remaining:
                # Bold **text**
                bold_match = re.search(r'\*\*(.+?)\*\*', remaining)
                italic_match = re.search(r'\*(.+?)\*', remaining)
                code_match = re.search(r'`(.+?)`', remaining)
                link_match = re.search(r'\[(.+?)\]\((.+?)\)', remaining)
                
                matches = []
                if bold_match:
                    matches.append(('bold', bold_match.start(), bold_match.end(), bold_match.group(1)))
                if italic_match and not bold_match:
                    matches.append(('italic', italic_match.start(), italic_match.end(), italic_match.group(1)))
                if code_match:
                    matches.append(('code', code_match.start(), code_match.end(), code_match.group(1)))
                if link_match:
                    matches.append(('link', link_match.start(), link_match.end(), link_match.group(1)))
                
                if matches:
                    # Find the earliest match
                    match_type, start, end, content_text = min(matches, key=lambda x: x[1])
                    
                    # Add text before match
                    if start > 0:
                        p.add_run(remaining[:start])
                    
                    # Add matched content
                    run = p.add_run(content_text)
                    if match_type == 'bold':
                        run.font.bold = True
                    elif match_type == 'italic':
                        run.font.italic = True
                    elif match_type == 'code':
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    elif match_type == 'link':
                        # Extract URL from link_match
                        remaining_match = re.search(r'\[(.+?)\]\((.+?)\)', remaining[start:])
                        if remaining_match:
                            url = remaining_match.group(2)
                            # Add hyperlink (simplified - just add as text with URL)
                            run.font.underline = True
                            run.font.color.rgb = RGBColor(0, 0, 255)
                    
                    remaining = remaining[end:]
                else:
                    # No more matches
                    if remaining:
                        p.add_run(remaining)
                    break
        else:
            # Empty line - add paragraph break
            if i > 0:
                doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"✓ Successfully converted to: {output_file}")
    print(f"  File size: {os.path.getsize(output_file) / 1024:.1f} KB")


if __name__ == '__main__':
    markdown_file = r'F:\hotel-booking-platform\backend-design-defense.md'
    output_file = r'F:\hotel-booking-platform\backend-design-defense.docx'
    
    if not os.path.exists(markdown_file):
        print(f"Error: File not found - {markdown_file}")
    else:
        md_to_docx(markdown_file, output_file)
